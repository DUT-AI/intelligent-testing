
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

C = {
    "RED": "C0392B", "ORANGE": "D35400", "GREEN": "1A7A4A",
    "PURPLE": "6C3483", "BLUE": "1A5276", "GRAY_BG": "F2F3F4",
    "LIGHT_BLUE": "D6EAF8", "LIGHT_GREEN": "D5F5E3", "LIGHT_RED": "FDEDEC",
    "LIGHT_ORG": "FDEBD0", "LIGHT_PUR": "E8DAEF", "LIGHT_GRAY": "F8F9FA",
    "WHITE": "FFFFFF", "DARK_TEXT": "1C2833",
}

def set_cell_bg(cell, color_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()

# Helpers
def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(C["DARK_TEXT"])

def h2(text, color=C["DARK_TEXT"]):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)

def h3(text, color=C["DARK_TEXT"]):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)

def para(runs):
    p = doc.add_paragraph()
    for text, bold, code, color in runs:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if bold: r.bold = True
        if code:
            r.font.name = 'Courier New'
            r.font.color.rgb = RGBColor.from_string("8B0000")
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

def bullet(runs, lvl=0):
    p = doc.add_paragraph(style='List Bullet' if lvl==0 else 'List Bullet 2')
    for text, bold, code, color in runs:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if bold: r.bold = True
        if code:
            r.font.name = 'Courier New'
            r.font.color.rgb = RGBColor.from_string("8B0000")
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

def notebox(label, text, bg, border_c):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + "\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(border_c)
    p.add_run(text)
    doc.add_paragraph()

def draw_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], C["BLUE"])
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
            hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        # support rows passed as either (row_data, bg) or as a flat list
        if isinstance(row, (list, tuple)) and len(row) == len(headers) + 1:
            row_data = list(row[:-1])
            bg = row[-1]
        elif isinstance(row, (list, tuple)) and len(row) == 2 and isinstance(row[0], (list, tuple)):
            row_data, bg = row
        else:
            row_data = row
            bg = None
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            if bg and bg != C["WHITE"]:
                set_cell_bg(row_cells[i], bg)
    doc.add_paragraph()

# TITLE PAGE
t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run("Content Database")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor.from_string(C["BLUE"])

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Feature Engineering cho Dataset C++ (Code-Only)")
r2.font.size = Pt(14)
r2.italic = True
r2.font.color.rgb = RGBColor.from_string("555555")

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Dành riêng cho câu hỏi dạng output prediction / code tracing")
r3.font.color.rgb = RGBColor.from_string("888888")

notebox("⚠ Phạm vi áp dụng", "Tài liệu này được viết lại hoàn toàn để phù hợp với dataset chỉ chứa câu hỏi code C++. Mỗi câu hỏi bao gồm một đoạn code và 5 phương án trả lời (output, 'Compile Error', 'Runtime Error', 'Another Answer'). Các đặc trưng về ngôn ngữ tự nhiên (văn phong, cú pháp tiếng Việt) được thay thế hoàn toàn bằng các đặc trưng phân tích cấu trúc code.", C["LIGHT_RED"], C["RED"])

# PART 1
h1("PHẦN 1 — ĐẶC TRƯNG THÔ (Surface Features)")
para([("Mục tiêu: ", True, False, C["BLUE"]), ("Trích xuất bằng regex / AST tĩnh, không cần LLM, chi phí gần bằng 0.", False, False, None)])

h2("1A. Đặc trưng từ vựng code (Code Lexical Features)", C["GREEN"])
h3("Số token câu hỏi (Question Token Count)")
bullet([("Cách tính: ", True, False, None), ("Đếm tổng số token trong phần mô tả câu hỏi (tiếng Việt) — không bao gồm code snippet.", False, False, None)])
bullet([("Phương pháp: ", True, False, None), ("Tokenize bằng underthesea hoặc PhoBERT tokenizer.", False, False, None)])
bullet([("Mục đích: ", True, False, None), ("Đo khối lượng ngôn ngữ tự nhiên sinh viên phải đọc trước khi nhìn vào code.", False, False, None)])
bullet([("Ghi chú thực tế: ", True, False, None), ("Với dataset hiện tại (câu hỏi kiểu 'Kết quả thực hiện đoạn chương trình sau là gì?'), giá trị thường rất thấp và ổn định. Feature này ít discriminative.", False, False, None)])

h3("Số dòng code (Code Line Count)")
bullet([("Cách tính: ", True, False, None), ("Đếm số dòng trong code snippet, bao gồm cả dòng trống và comment.", False, False, None)])
bullet([("Công thức: ", True, False, None), ("line_count = len(code.strip().split('\n'))", False, True, None)])
bullet([("Mục đích: ", True, False, None), ("Đo 'khối lượng vật lý' của đoạn code cần trace.", False, False, None)])
bullet([("Tác dụng: ", True, False, None), ("Code dài hơn → sinh viên cần duy trì nhiều trạng thái biến trong working memory hơn → độ khó tăng.", False, False, None)])

h3("Mật độ từ khóa C++ (C++ Keyword Density)")
bullet([("Cách tính: ", True, False, None), ("Tỷ lệ từ khóa C++ xuất hiện trên tổng số token trong code.", False, False, None)])
bullet([("Công thức: ", True, False, None), ("keyword_density = count(keywords_found) / total_code_tokens", False, True, None)])
bullet([("Danh sách từ khóa phân tầng:", True, False, None)])
bullet([("Tầng 1 – Cơ bản: ", True, False, None), ("int, float, double, char, bool, void, if, else, for, while, do, return, cout, cin", False, True, None)], 1)
bullet([("Tầng 2 – Trung cấp: ", True, False, None), ("class, struct, enum, namespace, const, static, reference (&), pointer (*)", False, True, None)], 1)
bullet([("Tầng 3 – Nâng cao: ", True, False, None), ("virtual, override, template, friend, new, delete, dynamic_cast, typeid", False, True, None)], 1)
bullet([("Tác dụng: ", True, False, None), ("Keyword density cao → câu hỏi tập trung nhiều cơ chế C++ → nhận thức phức tạp hơn.", False, False, None)])

h3("Số identifier người dùng định nghĩa (User-Defined Identifiers)")
bullet([("Cách tính: ", True, False, None), ("Đếm các tên biến, hàm, lớp do lập trình viên đặt (không phải từ khóa C++ hoặc standard library).", False, False, None)])
bullet([("Phương pháp: ", True, False, None), ("re.findall(r'\\b[a-zA-Z_][a-zA-Z0-9_]*\\b', code)", False, True, None), (" sau đó loại bỏ từ khóa C++ và STL.", False, False, None)])
bullet([("Tác dụng: ", True, False, None), ("Nhiều identifier → sinh viên phải track nhiều biến cùng lúc → tải working memory cao hơn.", False, False, None)])

h2("1B. Đặc trưng cú pháp code (Code Syntactic Features)", C["GREEN"])
h3("Độ sâu lồng ghép tối đa (Max Nesting Depth)")
bullet([("Cách tính: ", True, False, None), ("Đếm số cặp dấu ngoặc nhọn {} lồng nhau sâu nhất trong code.", False, False, None)])
bullet([("Phương pháp:", True, False, None)])
bullet([('depth = 0; max_depth = 0\nfor ch in code:\n    if ch == "{": depth += 1; max_depth = max(max_depth, depth)\n    elif ch == "}": depth -= 1', False, True, None)], 1)
bullet([("Tác dụng: ", True, False, None), ("Mỗi cấp lồng ghép thêm 1 scope context mà sinh viên phải giữ trong đầu. Độ sâu >= 4 thường rất khó trace.", False, False, None)])

h3("Độ phức tạp cấu trúc điều khiển (Control Flow Complexity)")
bullet([("Cách tính: ", True, False, None), ("Đếm tổng số cấu trúc điều khiển trong code.", False, False, None)])
bullet([("Công thức: ", True, False, None), ("cf_score = count(if/else) + count(for) * 1.5 + count(while/do-while) * 1.5 + count(switch) * 1.2", False, True, None)])
bullet([("Lý do trọng số: ", True, False, None), ("Vòng lặp (for/while) yêu cầu sinh viên track iteration state, khó hơn if-else đơn giản.", False, False, None)])

h3("Số lượng toán tử đặc thù C++ (C++-Specific Operator Count)")
bullet([("Cách tính: ", True, False, None), ("Đếm tần suất các toán tử gây tải nhận thức cao, tính tỷ lệ từng loại trên tổng code tokens.", False, False, None)])

draw_table(["Toán tử", "Mẫu regex", "Trọng số", "Lý do khó"], [
    ["* (pointer/deref)", "r'(?<![a-zA-Z])\*'", "1.5", "Nhầm lẫn pointer vs value", C["LIGHT_RED"]],
    ["& (ref/address)", "r'&(?!=)'", "1.5", "Nhầm tham chiếu vs địa chỉ", C["WHITE"]],
    ["-> (member ptr)", "r'->'", "2.0", "Yêu cầu hiểu pointer + struct", C["LIGHT_RED"]],
    [":: (scope)", "r'::'", "1.2", "Cần hiểu namespace/class scope", C["WHITE"]],
    ["++ / -- (pre/post)", "r'\+\+|--'", "1.8", "Pre vs post dễ gây lỗi nhất", C["LIGHT_RED"]],
    ["= vs == confusion", "r'=[^=]' vs r'=='", "0.8", "Lỗi cú pháp bẫy phổ biến", C["WHITE"]]
])
bullet([("Vector đầu ra: ", True, False, None), ("[d_ptr, d_ref, d_arrow, d_scope, d_incr, d_assign]", False, True, None), (" — 6 chiều tỷ lệ xuất hiện", False, False, None)])

h2("1C. Đặc trưng cấu trúc (Code Structural Features)", C["GREEN"])
h3("Số lớp / struct được định nghĩa (Class/Struct Count)")
bullet([("Cách tính: ", True, False, None), ("re.findall(r'\\b(class|struct)\\s+\\w+', code)", False, True, None)])
h3("Chỉ số quan hệ OOP (OOP Relationship Index)")
bullet([("Thành phần:", True, False, None)])
bullet([("N_class: ", True, False, None), ("Số lượng class/struct được khai báo.", False, False, None)], 1)
bullet([("D_inherit: ", True, False, None), ("Độ sâu kế thừa tối đa.", False, False, None)], 1)
bullet([("R_count: ", True, False, None), ("Số mối quan hệ tổng cộng.", False, False, None)], 1)
bullet([("Vector: ", True, False, None), ("[N_class, D_inherit, R_count]", False, True, None)])

h3("Phân loại vùng nhớ (Memory Zone Classification)")
bullet([("Vector nhị phân: ", True, False, None), ("[is_stack, is_heap, is_static, is_global]", False, True, None)])

h3("Phân loại kiểu dữ liệu phức tạp nhất (Max Type Complexity)")
draw_table(["Mức", "Kiểu dữ liệu", "Ví dụ"], [
    ["1", "Kiểu nguyên thủy", "int, float, char, bool", C["LIGHT_GREEN"]],
    ["2", "Kiểu dẫn xuất cơ bản", "int[], struct S, enum E", C["WHITE"]],
    ["3", "Pointer / Reference", "int*, int&, SinhVien*", C["LIGHT_ORG"]],
    ["4", "Pointer to ptr / Array of ptr", "int**, char*[]", C["LIGHT_RED"]],
    ["5", "Template / Smart pointer", "vector<T>, shared_ptr<T>", C["LIGHT_RED"]]
])

h2("1D. Đặc trưng phương án trả lời (Option Features)", C["GREEN"])
h3("Variance độ dài phương án (Option Length Variance)")
h3("Tỷ lệ phương án đặc biệt (Special Option Ratio)")
h3("Độ tương đồng giữa các phương án số (Numeric Option Similarity)")

h1("PHẦN 2 — ĐẶC TRƯNG ẨN (LLM-Extracted Features)")
notebox("Quy trình chuẩn cho mọi LLM feature", "1. Chạy prompt 3 lần với temperature khác nhau...\n2. Với giá trị số: lấy mean...\n3. Nếu 3 lần cho kết quả chênh lệch lớn...\n4. Luôn yêu cầu LLM suy luận trong <thinking>...", C["LIGHT_BLUE"], "2471A3")

h2("2A. Độ phức tạp quy trình (Procedural Complexity)", C["BLUE"])
h3("Số bước trace thực thi nguyên tử (Atomic Execution Steps)")
draw_table(["Loại bước", "Định nghĩa", "Ví dụ trong code"], [
    ["Identity", "Xác định loại thực thể", "int *ptr -> ptr là con trỏ kiểu int", C["LIGHT_GREEN"]],
    ["Simulation", "1 lần thay đổi giá trị", "var_x=2, ++var_x -> var_x=3", C["WHITE"]],
    ["Retrieval", "Gọi lại quy tắc C++", "ref-- là post-decrement", C["LIGHT_BLUE"]],
    ["Constraint Check", "Kiểm tra scope / type / const", "ref_x = var_x -> alias", C["WHITE"]]
])
notebox("Prompt mẫu", "SYSTEM: Bạn là chuyên gia phân tích độ khó câu hỏi lập trình C++...\n[ATOMIC_STEP]...\nTrọng số: Identityx1.0, Retrievalx1.5, Simulationx1.0, Constraintx1.2...", "EBF5FB", "2471A3")

h2("2B. Mức độ Bloom (Cognitive Level)", C["BLUE"])
draw_table(["Mức", "Tên", "Hành vi cần thiết trong C++ code"], [
    ["1", "Remember", "Nhận biết cú pháp", C["LIGHT_GREEN"]],
    ["2", "Understand", "Hiểu nghĩa của đoạn code ngắn", C["LIGHT_GREEN"]],
    ["3", "Apply", "Áp dụng quy tắc C++ vào code cụ thể", C["LIGHT_BLUE"]],
    ["4", "Analyze", "Trace nhiều biến/tham chiếu tương tác", C["LIGHT_ORG"]],
    ["5", "Evaluate", "Phán đoán hành vi phức tạp", C["LIGHT_RED"]],
    ["6", "Create", "Phát hiện UB / edge case", C["LIGHT_RED"]]
])
notebox("Lưu ý quan trọng", "Số bước (N) và Bloom level (B) KHÔNG tuyệt đối tỷ lệ thuận...", C["LIGHT_ORG"], C["ORANGE"])

h2("2C. Vector kỹ năng (Skill Identification Vector)", C["BLUE"])
draw_table(["#", "Skill ID", "Mô tả"], [
    ["1", "SK_REF", "Reference variable", C["LIGHT_BLUE"]],
    ["2", "SK_PTR", "Pointer declaration", C["WHITE"]],
    ["3", "SK_PRE_POST", "Pre vs post side effect", C["LIGHT_BLUE"]],
    ["4", "SK_ENUM", "Enum definition", C["WHITE"]],
    ["5", "SK_LOOP", "Loop execution trace", C["LIGHT_BLUE"]],
    ["6", "SK_SCOPE", "Variable scope", C["WHITE"]],
    ["7", "SK_COUT_ORDER", "cout evaluation order", C["LIGHT_BLUE"]],
    ["8", "SK_OOP_SINGLE", "Single class", C["WHITE"]],
    ["9", "SK_OOP_INHERIT", "Inheritance", C["LIGHT_RED"]],
    ["10", "SK_VIRTUAL", "Virtual function", C["LIGHT_RED"]],
    ["11", "SK_MEMORY", "Dynamic allocation", C["LIGHT_RED"]],
    ["12", "SK_OPERATOR", "Operator overloading", C["WHITE"]]
])

h2("2D. Phân tích quan niệm sai lầm (Misconception Analysis)", C["BLUE"])
h3("Phân loại nhóm lỗi đặc thù cho code C++")
draw_table(["Loại lỗi", "Ví dụ cụ thể", "Trọng số P"], [
    ["Conceptual", "Nhầm & là địa chỉ vs tham chiếu", "0.8", C["LIGHT_RED"]],
    ["Procedural", "Quên delete[]", "0.5", C["LIGHT_ORG"]],
    ["Side-effect trap", "++x vs x++ trong cout", "1.0", C["LIGHT_RED"]],
    ["Syntactic trap", "= thay vì ==", "0.6", C["LIGHT_ORG"]],
    ["Lack of sense", "Không biết enum", "1.2", C["LIGHT_RED"]]
])

h2("2E. Độ khó của phương án nhiễu (Distractor Plausibility)", C["BLUE"])

h1("PHẦN 3 — VECTOR ĐẶC TRƯNG TỔNG HỢP (Feature Vector Summary)")
draw_table(["#", "Nhóm đặc trưng", "Số chiều", "Ký hiệu"], [
    ["1-14", "Surface Features (như L_lines, S_nest...)", "25", "L_, S_, T_, O_", C["LIGHT_GREEN"]],
    ["15-24", "LLM Features (như H_N, H_B, H_S...)", "~25", "H_", C["LIGHT_BLUE"]],
    ["25", "Code text embedding (ModernBERT)", "768", "E_code", C["LIGHT_PUR"]],
    ["—", "TỔNG CỘNG", "~806 chiều", "—", C["GRAY_BG"]]
])
notebox("Thứ tự ưu tiên triển khai", "Phase 1: Text embedding\nPhase 2: Surface features\nPhase 3: LLM features", C["LIGHT_GREEN"], "1A7A4A")

h1("PHẦN 4 — VÍ DỤ ÁP DỤNG (Question ID 7 & 10)")
notebox("Insight từ ví dụ", "Câu 7 và 10 đều có SK_REF + SK_PRE_POST, nhưng câu 10 khó hơn nhiều vì có loop và var_x++...", C["LIGHT_PUR"], C["PURPLE"])

h1("PHẦN 5 — ĐẶC TRƯNG BỎ ĐI SO VỚI CONTENT DATABASE GỐC")
draw_table(["Feature gốc", "Lý do bỏ / thay thế", "Thay bằng"], [
    ["Word Count", "Không phân biệt", "Code line count", C["LIGHT_RED"]],
    ["Syntactic Complexity", "Không có trong code", "Max nesting depth", C["LIGHT_RED"]]
])

doc.save('artifacts/Content_Database_v2_CodeOnly_Full.docx')
